from pathlib import Path
from typing import Dict, Any, Optional
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting transcription results to various formats"""

    def __init__(self):
        self.templates_dir = Path(__file__).parent.parent / "templates" / "word"
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def export_to_word(
        self,
        transcription_data: Dict[str, Any],
        template_path: Optional[Path] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """
        Export transcription results to a Word document

        Args:
            transcription_data: Transcription results with segments and speaker info
            template_path: Optional path to a Word template file
            output_path: Optional output path for the generated document

        Returns:
            Path to the generated Word document
        """
        try:
            # Use template if provided, otherwise create new document
            if template_path and template_path.exists():
                logger.info(f"Using template: {template_path}")
                doc = self._populate_template(template_path, transcription_data)
            else:
                logger.info("Creating document from default format")
                doc = self._create_default_document(transcription_data)

            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = Path(f"transcription_{timestamp}.docx")

            # Save the document
            doc.save(str(output_path))
            logger.info(f"Word document saved to: {output_path}")

            return output_path

        except Exception as e:
            logger.error(f"Failed to export to Word: {e}")
            raise RuntimeError(f"Word export failed: {str(e)}")

    def _create_default_document(self, data: Dict[str, Any]) -> Document:
        """Create a Word document with default formatting"""
        doc = Document()

        # Add title
        title = doc.add_heading("Transcription Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Check if this is transcription-only mode (single speaker named "Speaker 1")
        is_transcription_only = (
            data.get('num_speakers', 0) == 1 and
            list(data.get('speakers', {}).values()) == ["Speaker 1"]
        )

        # Add metadata section
        doc.add_heading("Metadata", level=1)

        # Adjust metadata rows based on mode
        if is_transcription_only:
            metadata_table = doc.add_table(rows=4, cols=2)
            metadata_table.style = 'Light Grid Accent 1'
            metadata_rows = [
                ("Export Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                ("Duration", f"{data.get('duration', 0):.2f} seconds"),
                ("Language", data.get('language', 'Unknown').upper()),
                ("Total Segments", str(len(data.get('segments', []))))
            ]
        else:
            metadata_table = doc.add_table(rows=5, cols=2)
            metadata_table.style = 'Light Grid Accent 1'
            metadata_rows = [
                ("Export Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                ("Duration", f"{data.get('duration', 0):.2f} seconds"),
                ("Number of Speakers", str(data.get('num_speakers', 0))),
                ("Language", data.get('language', 'Unknown').upper()),
                ("Total Segments", str(len(data.get('segments', []))))
            ]

        for i, (key, value) in enumerate(metadata_rows):
            metadata_table.rows[i].cells[0].text = key
            metadata_table.rows[i].cells[1].text = value
            # Make key column bold
            metadata_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # Add spacing

        # Add transcription section
        doc.add_heading("Transcription", level=1)

        segments = data.get('segments', [])

        if not segments:
            doc.add_paragraph("No transcription segments available.")
        else:
            if is_transcription_only:
                # Transcription-only mode: Just show text without speaker/timestamp
                for segment in segments:
                    text = segment.get('text', '').strip()
                    text_para = doc.add_paragraph(text)
                    text_para.paragraph_format.space_after = Pt(12)
            else:
                # Standard mode: Show speaker and timestamp
                for segment in segments:
                    # Create speaker header
                    speaker = segment.get('speaker', 'Unknown Speaker')
                    start_time = segment.get('start', 0)
                    end_time = segment.get('end', 0)

                    # Format timestamp as MM:SS
                    start_str = self._format_timestamp(start_time)
                    end_str = self._format_timestamp(end_time)

                    # Add speaker and timestamp
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"{speaker}")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(11)
                    speaker_run.font.color.rgb = RGBColor(0, 70, 140)

                    timestamp_run = speaker_para.add_run(f"  [{start_str} - {end_str}]")
                    timestamp_run.font.size = Pt(10)
                    timestamp_run.font.color.rgb = RGBColor(100, 100, 100)
                    timestamp_run.italic = True

                    # Add transcription text
                    text = segment.get('text', '').strip()
                    text_para = doc.add_paragraph(text)
                    text_para.paragraph_format.left_indent = Inches(0.25)
                    text_para.paragraph_format.space_after = Pt(12)

        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated by Audio Scribe AI")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer_run.italic = True

        return doc

    def _populate_template(self, template_path: Path, data: Dict[str, Any]) -> Document:
        """
        Populate a Word template with transcription data

        This method looks for placeholders in the template and replaces them:
        - {{title}} - Document title
        - {{date}} - Export date
        - {{duration}} - Audio duration
        - {{num_speakers}} - Number of speakers
        - {{language}} - Detected language
        - {{TRANSCRIPTION}} - Special marker where transcription segments will be inserted
        """
        doc = Document(str(template_path))

        # Prepare replacement values
        replacements = {
            '{{title}}': 'Transcription Report',
            '{{date}}': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            '{{duration}}': f"{data.get('duration', 0):.2f} seconds",
            '{{num_speakers}}': str(data.get('num_speakers', 0)),
            '{{language}}': data.get('language', 'Unknown').upper(),
            '{{total_segments}}': str(len(data.get('segments', [])))
        }

        # Replace placeholders in paragraphs
        transcription_marker_found = False
        transcription_insert_index = None

        for i, paragraph in enumerate(doc.paragraphs):
            # Check for transcription marker
            if '{{TRANSCRIPTION}}' in paragraph.text:
                transcription_marker_found = True
                transcription_insert_index = i
                # Clear the marker paragraph
                paragraph.clear()
                continue

            # Replace other placeholders
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    # Replace inline while preserving formatting
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)

        # Insert transcription segments
        if transcription_marker_found and transcription_insert_index is not None:
            segments = data.get('segments', [])

            # Check if this is transcription-only mode
            is_transcription_only = (
                data.get('num_speakers', 0) == 1 and
                list(data.get('speakers', {}).values()) == ["Speaker 1"]
            )

            # Insert segments at the marker position
            for segment in segments:
                text = segment.get('text', '').strip()

                if is_transcription_only:
                    # Transcription-only mode: Just insert text without speaker/timestamp
                    text_para = doc.add_paragraph()
                    doc._element.body.insert(transcription_insert_index, text_para._element)

                    text_para = doc.paragraphs[transcription_insert_index]
                    text_para.add_run(text)
                    text_para.paragraph_format.space_after = Pt(12)

                    transcription_insert_index += 1
                else:
                    # Standard mode: Insert speaker, timestamp, and text
                    speaker = segment.get('speaker', 'Unknown Speaker')
                    start_time = segment.get('start', 0)
                    end_time = segment.get('end', 0)
                    start_str = self._format_timestamp(start_time)
                    end_str = self._format_timestamp(end_time)

                    # Insert speaker paragraph at index
                    speaker_para = doc.paragraphs[transcription_insert_index]._element
                    new_speaker_para = doc.add_paragraph()
                    doc._element.body.insert(transcription_insert_index, new_speaker_para._element)

                    # Format speaker paragraph
                    new_speaker_para = doc.paragraphs[transcription_insert_index]
                    speaker_run = new_speaker_para.add_run(f"{speaker}")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(11)
                    speaker_run.font.color.rgb = RGBColor(0, 70, 140)

                    timestamp_run = new_speaker_para.add_run(f"  [{start_str} - {end_str}]")
                    timestamp_run.font.size = Pt(10)
                    timestamp_run.font.color.rgb = RGBColor(100, 100, 100)
                    timestamp_run.italic = True

                    transcription_insert_index += 1

                    # Insert text paragraph
                    text_para = doc.add_paragraph()
                    doc._element.body.insert(transcription_insert_index, text_para._element)

                    text_para = doc.paragraphs[transcription_insert_index]
                    text_para.add_run(text)
                    text_para.paragraph_format.left_indent = Inches(0.25)
                    text_para.paragraph_format.space_after = Pt(12)

                    transcription_insert_index += 1

        return doc

    def _format_timestamp(self, seconds: float) -> str:
        """Format seconds as MM:SS or HH:MM:SS"""
        total_seconds = int(seconds)
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        secs = total_seconds % 60

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"

    def get_available_templates(self) -> list[Path]:
        """Get list of available Word templates"""
        if not self.templates_dir.exists():
            return []

        return list(self.templates_dir.glob("*.docx"))
